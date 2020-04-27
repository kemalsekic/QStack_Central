import docx
import os
import shelve
import pprint
import subprocess
import configStuff
import stagingFile

doc = docx.Document()
docName = "docStuff"

def addDetailToWord():
   #Create the Document and add headers
   doc.add_heading('Cubby - Staging Device', 0)
   doc.add_paragraph("Username: " + stagingFile.uName + "\t\t\tPassword: " + stagingFile.pWord)
   doc.add_paragraph("Account #: " + stagingFile.accNum + "\t\t\t\tSSN: " + stagingFile.ssn)
   doc.add_paragraph("Name: " + stagingFile.firstName + " " + stagingFile.lastName
   + "\t\t\t\t\tDOB: " + stagingFile.dobMonth + "/" + stagingFile.dobDay + "/" + stagingFile.dobYear)
   doc.add_paragraph("Address: " + stagingFile.address + "\t\t\t\t\tAddress Line 2: " + stagingFile.address2)
   doc.add_paragraph("City/State/Zip: " + stagingFile.city + "/" + stagingFile.state + "/" + stagingFile.zipCode
   + "\t\t\t\tMDN: " + stagingFile.mdn)
   doc.add_paragraph("Invoice #: " + stagingFile.invoiceNum + "\t\t\t\t\t\tOrder #: " + stagingFile.orderNum)

def addScreenShot(imageName, headingText):
    configStuff.count += 1
    doc.add_heading("Step " + str(configStuff.count) + " - " + headingText, 2)
    doc.add_picture(imageName + '.png', width=docx.shared.Inches(6.5), height=docx.shared.Inches(3.66))
    os.remove('C:/Users/P2938214/Documents/Chemal/QStack/' + imageName + '.png')
    doc.save(docName + '.docx')

