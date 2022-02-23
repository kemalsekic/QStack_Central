import docx
import os
import shelve
import pprint
import subprocess
import webbrowser
import config.configStuff as configStuff
import config.driverThings as driverThings
import deviceDetails.deviceDetail as deviceDetail
from os import system, name

driver = driverThings.Driver
doc = docx.Document()

def clearS():
    if name == 'nt': 
        _ = system('cls') 

def addScreenShot(imageName):
    configStuff.count += 1
    print(configStuff.count)
    driver.save_screenshot(imageName + str(configStuff.count) + ".png")
    doc.add_heading('QStack', 2)
    doc.add_picture(imageName + str(configStuff.count) + '.png', width=docx.shared.Inches(6), height=docx.shared.Inches(5))

#Create the Document and add headers
clearS()
docName = input("Name of document: ")
doc.add_heading('Cubby - Staging Device', 0)

#f = open("num.txt", "w+")

#Load site
driver.maximize_window()
driver.get('https://buy-uat.spectrummobile.com/')

#Wait for element to display then click it
driver.implicitly_wait(20)
byod = driver.find_element_by_link_text('Bring Your Device')
byod.click()

clearS()
input("Please wait for the screen to load.\nThen press Enter to continue..")
clearS()

#Finding elements on page
selectDevice = driverThings.Select(driver.find_element_by_id('devicetype'))
selectManuf = driverThings.Select(driver.find_element_by_id('manufacturer'))
deviceIMEI = driver.find_element_by_name('imei')
checkMyDeviceButton = driver.find_element_by_id('checkMyDevice')

selectDevice.select_by_visible_text('Phone')
selectManuf.select_by_visible_text('Apple')
deviceIMEI.send_keys(deviceDetail.IMEI)

#Wait for user to get through captcha
clearS()
input("Click the captcha then type any key to continue.")
addScreenShot("screenshot")

#Click the button to continue
checkMyDeviceButton.click()
driver.implicitly_wait(10)
addScreenShot("screenshot")

doc.save(docName + '.docx')

clearS()
print("\t\tStaging Complete.\nCheck your document: " + docName + ".docx")
print("\n\nCubby\n\t\tBy QStack\n\n\n\n")
#driver.quit()

