import docx
import os
import shelve
import pprint
import subprocess
import webbrowser
from selenium import webdriver

driver = webdriver.Chrome()
#type(browser)

driver.get('https://bmw-agility.ridgelinemobile.net/csr/f?p=2260:101:15851567913649')

UNField = driver.find_element_by_id('P101_USERNAME')
passField = driver.find_element_by_id('P101_PASSWORD')
signInButton = driver.find_element_by_id('button-login')

UNField.send_keys('p2938214')
passField.send_keys('')
signInButton.click()

#webbrowser.open('https://bmw-agility.ridgelinemobile.net/csr/f?p=2260:101:15851567913649')

#shelfFile = shelve.open('mydata')
#username = input("username")
#password = input("password")
#cats = ['kitty', 'Ajla', 'Kremi']
#shelfFile.close()

doc = docx.Document('kemo.docx')
print(len(doc.paragraphs))

print(doc.paragraphs[1].text)
doc2 = docx.Document()
doc2.add_heading('Kemal Sekic', 0)
doc2.add_heading('QStack', 2)
#doc2.add_paragraph(cats[2])

doc2.add_picture('kem.png', width=docx.shared.Inches(6), height=docx.shared.Inches(5))
doc2.save('kemal.docx')
#os.startfile('C:/Users/P2938214/Documents/Chemal/QStack/kemal.docx')

#exec('SShots.py')
#here = input("here")